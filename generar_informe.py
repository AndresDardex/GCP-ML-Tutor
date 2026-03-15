"""
generar_informe.py
==================
Genera el informe Word del proyecto MentorML.
Uso: python generar_informe.py

Antes de ejecutar, toma los siguientes pantallazos y guárdalos en la
carpeta docs/img/ con estos nombres exactos:

  01_home.png          → Pantalla de inicio con los 3 botones
  02_flashcard_term.png → Flashcard mostrando un término (antes de responder)
  03_flashcard_answer.png → Área de texto con una respuesta escrita
  04_feedback_correct.png  → Resultado correcto (verde, puntuación alta)
  05_feedback_partial.png  → Resultado parcial (amarillo)
  06_feedback_incorrect.png → Resultado incorrecto (rojo)
  07_official_def.png  → Expander abierto mostrando la definición oficial
  08_chat_question.png → Chat con una pregunta escrita
  09_chat_answer.png   → Chat mostrando la respuesta del tutor
  10_progress.png      → Pantalla de progreso con métricas y barras
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalando python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG_DIR = os.path.join(os.path.dirname(__file__), "docs", "img")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "docs", "MentorML_Informe.docx")


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color and h.runs:
        h.runs[0].font.color.rgb = RGBColor(*color)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


def add_image_or_placeholder(doc, filename, caption, width=5.5):
    img_path = os.path.join(IMG_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ PANTALLAZO: {filename} ]")
        run.font.color.rgb = RGBColor(180, 0, 0)
        run.font.bold = True
        run.font.size = Pt(11)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F4F4F4")
    p._p.get_or_add_pPr().append(shading)
    return p


def build_document():
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MentorML")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Tutor Académico de Machine Learning")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(74, 144, 226)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Proyecto: Desarrollo de un Asistente Experto basado en RAG y Agentes\n").font.bold = True
    meta.add_run("Avance 1 — Diseño de Prompts, Few-Shot y Arquitectura RAG\n")
    meta.add_run("Sistema de IA local para aprendizaje de Machine Learning")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. DESCRIPCIÓN GENERAL
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "1. Descripción General", 1, (30, 58, 95))
    add_body(doc,
        "MentorML es un asistente tutor académico especializado en Machine Learning que opera "
        "completamente de forma local, sin enviar ningún dato a servicios externos. Usa el "
        "Glosario oficial de ML de Google (697 términos) como base de conocimientos y un modelo "
        "de lenguaje local (Ollama + llama3.2) para evaluar respuestas del estudiante, responder "
        "preguntas y hacer seguimiento de su progreso."
    )
    add_body(doc,
        "El enfoque es el de un Tutor Académico Personalizado, uno de los enfoques contemplados "
        "en el proyecto, basado en el glosario oficial de Machine Learning de Google como material "
        "de referencia."
    )

    # ══════════════════════════════════════════════════════════════
    # 2. PANTALLA DE INICIO
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "2. Pantalla de Inicio", 1, (30, 58, 95))
    add_body(doc,
        "Al abrir la aplicación en http://localhost:8501 se presenta la pantalla principal con "
        "tres modos de uso disponibles: Flashcards, Pregunta Libre y Mi Progreso. También indica "
        "cuántos términos del glosario están cargados y listos para practicar."
    )
    add_image_or_placeholder(doc, "01_home.png", "Figura 1 — Pantalla de inicio de MentorML con los tres modos disponibles")

    # ══════════════════════════════════════════════════════════════
    # 3. MODO FLASHCARDS
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "3. Modo Flashcards", 1, (30, 58, 95))
    add_body(doc,
        "El modo Flashcards es la funcionalidad principal del tutor. Presenta un término de ML "
        "al estudiante, quien debe escribir su propia definición. Luego el sistema RAG evalúa "
        "la respuesta usando el LLM local y devuelve feedback detallado."
    )

    add_heading(doc, "3.1 Presentación del Término", 2)
    add_body(doc,
        "El término se muestra en una tarjeta azul destacada. El sistema de selección implementa "
        "spaced repetition: prioriza con 45% de probabilidad términos nunca vistos, con 40% los "
        "que tienen puntuación menor a 60, y con 30% los que están por debajo de 80."
    )
    add_image_or_placeholder(doc, "02_flashcard_term.png", "Figura 2 — Tarjeta con un término de ML listo para evaluar")

    add_heading(doc, "3.2 Respuesta del Estudiante", 2)
    add_body(doc,
        "El estudiante escribe su definición en el área de texto. Puede optar por saltar el "
        "término si no lo conoce. Al presionar 'Evaluar respuesta', el sistema construye el "
        "prompt completo (System Prompt + Few-Shot + RAG context) y lo envía al LLM."
    )
    add_image_or_placeholder(doc, "03_flashcard_answer.png", "Figura 3 — Estudiante escribiendo su respuesta antes de evaluar")

    add_heading(doc, "3.3 Resultado: Respuesta Correcta", 2)
    add_body(doc,
        "Cuando la puntuación es ≥ 80, el feedback aparece con fondo verde. Se muestra la "
        "puntuación numérica, una barra de progreso visual, y el feedback motivador del tutor. "
        "Si hay conceptos que podrían complementar la respuesta, también se listan."
    )
    add_image_or_placeholder(doc, "04_feedback_correct.png", "Figura 4 — Feedback verde: respuesta correcta con puntuación alta")

    add_heading(doc, "3.4 Resultado: Respuesta Parcial", 2)
    add_body(doc,
        "Cuando la puntuación está entre 60 y 79, el feedback aparece en amarillo indicando "
        "que la respuesta está en el camino correcto pero le faltan elementos importantes. "
        "Los conceptos clave faltantes se listan para que el estudiante los refuerce."
    )
    add_image_or_placeholder(doc, "05_feedback_partial.png", "Figura 5 — Feedback amarillo: respuesta parcialmente correcta")

    add_heading(doc, "3.5 Resultado: Respuesta Incorrecta", 2)
    add_body(doc,
        "Cuando la puntuación es menor a 60, el feedback aparece en rojo. El tutor explica "
        "qué es incorrecto o insuficiente de manera pedagógica y constructiva, sin desanimar "
        "al estudiante."
    )
    add_image_or_placeholder(doc, "06_feedback_incorrect.png", "Figura 6 — Feedback rojo: respuesta incorrecta o muy incompleta")

    add_heading(doc, "3.6 Definición Oficial del Glosario", 2)
    add_body(doc,
        "Después de recibir el feedback, el estudiante puede expandir un panel para ver la "
        "definición oficial tal como aparece en el Glosario de ML de Google, con todos sus "
        "párrafos y detalles."
    )
    add_image_or_placeholder(doc, "07_official_def.png", "Figura 7 — Definición oficial del glosario expandida")

    # ══════════════════════════════════════════════════════════════
    # 4. MODO PREGUNTA LIBRE (CHAT)
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "4. Modo Pregunta Libre (Chat)", 1, (30, 58, 95))
    add_body(doc,
        "En este modo el estudiante puede hacer cualquier pregunta sobre Machine Learning en "
        "lenguaje natural. El sistema RAG busca los 4 términos más relevantes del glosario y "
        "los inyecta como contexto antes de que el LLM genere la respuesta."
    )

    add_heading(doc, "4.1 Formulación de la Pregunta", 2)
    add_body(doc,
        "El chat acepta preguntas en español en lenguaje natural. Por ejemplo: '¿Cuál es la "
        "diferencia entre overfitting y underfitting?' o '¿Cómo funciona el backpropagation?'"
    )
    add_image_or_placeholder(doc, "08_chat_question.png", "Figura 8 — Pregunta libre escrita en el chat")

    add_heading(doc, "4.2 Respuesta del Tutor", 2)
    add_body(doc,
        "El tutor responde basándose en los términos recuperados del glosario. Las respuestas "
        "son precisas, en español, y mencionan términos relacionados cuando enriquece la "
        "explicación."
    )
    add_image_or_placeholder(doc, "09_chat_answer.png", "Figura 9 — Respuesta del tutor fundamentada en el glosario")

    # ══════════════════════════════════════════════════════════════
    # 5. MI PROGRESO
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "5. Pantalla de Progreso", 1, (30, 58, 95))
    add_body(doc,
        "La pantalla de progreso muestra un dashboard completo con el avance del estudiante: "
        "total de términos practicados, dominados (≥80 puntos), en repaso (<60 puntos) y el "
        "promedio general. También lista los términos específicos que necesitan atención."
    )
    add_image_or_placeholder(doc, "10_progress.png", "Figura 10 — Dashboard de progreso del estudiante")

    # ══════════════════════════════════════════════════════════════
    # 6. ARQUITECTURA RAG
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "6. Arquitectura RAG", 1, (30, 58, 95))
    add_body(doc,
        "RAG (Retrieval-Augmented Generation) es la técnica central del sistema. En lugar de "
        "depender únicamente del conocimiento del LLM, se recupera información relevante de la "
        "base de datos vectorial antes de generar cada respuesta."
    )

    add_heading(doc, "6.1 Flujo de Datos", 2)
    add_body(doc, "El proceso completo funciona así:")
    steps = [
        "1. El usuario escribe una respuesta o pregunta",
        "2. retrieve_context() convierte la consulta en un vector usando sentence-transformers",
        "3. ChromaDB busca los 3-4 documentos más similares por similitud coseno",
        "4. Los términos recuperados se inyectan en el prompt como <contexto_glosario>",
        "5. El prompt completo (System + Few-Shot + RAG + pregunta) se envía a Ollama",
        "6. Ollama genera la respuesta basada en el contexto real del glosario",
        "7. _parse_json_response() extrae y valida el JSON de la respuesta",
    ]
    for step in steps:
        add_bullet(doc, step)

    add_heading(doc, "6.2 Base de Datos Vectorial", 2)
    add_body(doc,
        "ChromaDB almacena 697 documentos, uno por cada término del glosario. Cada documento "
        "tiene el formato 'Término: X\\nDefinición: Y' y su vector de embedding de 384 dimensiones "
        "generado por el modelo paraphrase-multilingual-MiniLM-L12-v2. La búsqueda usa "
        "similitud coseno para encontrar los términos semánticamente más cercanos a la consulta."
    )

    # ══════════════════════════════════════════════════════════════
    # 7. DISEÑO DE PROMPTS (AVANCE 1)
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "7. Diseño de Prompts — Avance 1", 1, (30, 58, 95))

    add_heading(doc, "7.1 System Prompt con XML Tags", 2)
    add_body(doc,
        "Los system prompts usan XML tags como delimitadores para separar claramente el rol "
        "del asistente, las reglas de comportamiento y el contexto inyectado por RAG:"
    )
    add_code_block(doc,
        "<rol>\n"
        "Eres MentorML, un tutor experto en Machine Learning...\n"
        "</rol>\n\n"
        "<reglas>\n"
        "- Evalúa la respuesta del estudiante comparándola con la definición oficial\n"
        "- SIEMPRE responde en formato JSON válido, sin texto adicional fuera del JSON\n"
        "</reglas>\n\n"
        "<contexto_glosario>\n"
        "{context}   ← Términos recuperados por RAG\n"
        "</contexto_glosario>"
    )

    add_heading(doc, "7.2 Few-Shot Prompting", 2)
    add_body(doc,
        "Se incluyen 3 ejemplos de evaluación para guiar al modelo hacia el formato JSON "
        "esperado y el tono pedagógico deseado: un ejemplo de respuesta excelente (95 puntos), "
        "uno parcial (75 puntos) y uno incorrecto (20 puntos):"
    )
    add_code_block(doc,
        "<ejemplos_de_evaluacion>\n"
        "  <ejemplo_1>\n"
        '    <término>Overfitting</término>\n'
        "    <respuesta_estudiante>cuando el modelo aprende demasiado...</respuesta_estudiante>\n"
        "    <evaluacion_esperada>\n"
        "    {\n"
        '      "puntuacion": 75,\n'
        '      "correcto": true,\n'
        '      "feedback": "¡Bien! Captaste la idea principal...",\n'
        '      "conceptos_clave_faltantes": ["generalización", "ruido"]\n'
        "    }\n"
        "    </evaluacion_esperada>\n"
        "  </ejemplo_1>\n"
        "  ... (3 ejemplos en total)\n"
        "</ejemplos_de_evaluacion>"
    )

    add_heading(doc, "7.3 Formato de Salida JSON", 2)
    add_body(doc,
        "El prompt de evaluación especifica el esquema JSON exacto que debe retornar el modelo, "
        "lo que permite parsear la respuesta programáticamente y mostrarla de forma estructurada "
        "en la interfaz:"
    )
    add_code_block(doc,
        "{\n"
        '  "puntuacion": <número 0-100>,\n'
        '  "correcto": <true si puntuacion >= 60, false si no>,\n'
        '  "feedback": "<explicación motivadora de 2-3 oraciones>",\n'
        '  "conceptos_clave_faltantes": ["<concepto1>", "<concepto2>"]\n'
        "}"
    )

    add_heading(doc, "7.4 Estrategia de Delimitadores", 2)
    add_body(doc, "El proyecto usa dos estrategias de delimitadores complementarias:")
    add_bullet(doc,
        "XML tags (<rol>, <reglas>, <contexto_glosario>, <evaluacion_actual>, <término>): "
        "separan las secciones semánticas del prompt y son interpretados consistentemente por los LLMs"
    )
    add_bullet(doc,
        "Triple comillas Python (\"\"\"...\"\"\"): delimitan los strings multi-línea en el código "
        "fuente, facilitando la lectura y mantenimiento de los prompts"
    )

    # ══════════════════════════════════════════════════════════════
    # 8. STACK TECNOLÓGICO
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "8. Stack Tecnológico", 1, (30, 58, 95))

    techs = [
        ("Streamlit ≥1.31", "Interfaz web interactiva en Python, sin necesidad de HTML/JS"),
        ("Ollama + llama3.2", "LLM local, sin API externa, temperatura 0.3 para respuestas consistentes"),
        ("sentence-transformers", "Generación de embeddings semánticos multilingüe (paraphrase-multilingual-MiniLM-L12-v2)"),
        ("ChromaDB ≥0.5", "Base de datos vectorial local con búsqueda por similitud coseno"),
        ("LangChain ≥0.2", "Framework de orquestación para cadenas LLM"),
        ("BeautifulSoup4", "Extracción y parseo del glosario HTML de Google"),
        ("JSON local", "Persistencia del progreso del estudiante sin base de datos externa"),
    ]
    for tech, desc in techs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        run_bold = p.add_run(f"{tech}: ")
        run_bold.bold = True
        run_bold.font.color.rgb = RGBColor(30, 58, 95)
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════
    # CIERRE
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "9. Conclusión", 1, (30, 58, 95))
    add_body(doc,
        "MentorML demuestra cómo construir un sistema de IA educativo robusto y privado usando "
        "únicamente herramientas de código abierto. La combinación de RAG con un LLM local "
        "permite respuestas fundamentadas en una fuente de conocimiento verificable, mientras "
        "que el diseño cuidadoso de prompts (System Prompts, Few-Shot, XML tags y formato JSON) "
        "garantiza respuestas consistentes, estructuradas y pedagógicamente apropiadas."
    )
    add_body(doc,
        "El sistema corre completamente en la máquina del usuario, preservando la privacidad "
        "de los datos del estudiante y funcionando sin conexión a internet después del setup "
        "inicial."
    )

    return doc


def main():
    os.makedirs(os.path.join(os.path.dirname(__file__), "docs", "img"), exist_ok=True)
    print("Generando informe Word...")
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"\nInforme guardado en: {OUTPUT_PATH}")
    print("\n" + "="*60)
    print("PANTALLAZOS NECESARIOS")
    print("="*60)
    print("Guarda las imágenes en: docs/img/")
    print()
    screenshots = [
        ("01_home.png",              "Pantalla de inicio con los 3 botones"),
        ("02_flashcard_term.png",    "Flashcard mostrando un término (antes de responder)"),
        ("03_flashcard_answer.png",  "Área de texto con una respuesta escrita"),
        ("04_feedback_correct.png",  "Resultado correcto (verde, puntuación ≥80)"),
        ("05_feedback_partial.png",  "Resultado parcial (amarillo, 60-79)"),
        ("06_feedback_incorrect.png","Resultado incorrecto (rojo, <60)"),
        ("07_official_def.png",      "Expander abierto con la definición oficial"),
        ("08_chat_question.png",     "Chat con una pregunta escrita"),
        ("09_chat_answer.png",       "Chat mostrando la respuesta del tutor"),
        ("10_progress.png",          "Pantalla de progreso con métricas"),
    ]
    for filename, desc in screenshots:
        print(f"  {filename:35s} → {desc}")
    print()
    print("Después de guardar las imágenes, ejecuta este script de nuevo")
    print("para que queden embebidas en el documento Word.")


if __name__ == "__main__":
    main()
